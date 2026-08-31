#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
export.py — 把 Markdown 导出成 PDF 或 Word，带 360 品牌样式

用法:
    python export.py "输入.md" --to pdf              # 导出 PDF
    python export.py "输入.md" --to word             # 导出 Word
    python export.py "输入.md" --to pdf "输出.pdf"    # 指定输出路径

依赖:
    pip install markdown python-docx
    Edge/Chrome (PDF 用，Windows 自带 Edge)

设计要点(踩过的坑都规避了):
- 中文路径在某些 Windows Python 环境下 argv/glob 都会 mojibake：
  统一先把源文件复制到英文临时目录处理，产物再复制回目标路径
- CSS 用三引号常量，不做字符串拼接转义
- PDF: Edge headless --print-to-pdf, 用 Path.as_uri() 处理 URI
- Word: python-docx, 表格蓝底白字、图片自动嵌入、中文字体设置
- 标题无下划线横杠，hr 隐藏
"""
import sys
import os
import re
import time
import shutil
import base64
import tempfile
import subprocess
import argparse
from pathlib import Path

# ============ 360 品牌样式 (PDF 用) ============
CSS = """
@page { size: A4; margin: 2cm; }
body { font-family: "Microsoft YaHei", "微软雅黑", sans-serif; font-size: 14px; line-height: 1.8; color: #1a1a2e; }
h1 { color: #0721A8; font-size: 26px; margin-top: 30px; }
h2 { color: #0721A8; font-size: 22px; margin-top: 28px; }
h3 { color: #0D47A1; font-size: 18px; margin-top: 22px; }
h4 { color: #0D47A1; font-size: 16px; margin-top: 18px; }
table { border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 13px; }
th { background: #0721A8; color: #fff; padding: 8px 10px; text-align: left; font-weight: bold; }
td { border: 1px solid #ddd; padding: 7px 10px; }
tr:nth-child(even) { background: #f5f7fa; }
blockquote { border-left: 4px solid #0721A8; background: #e8eef6; padding: 10px 16px; margin: 12px 0; }
code { background: #f0f0f0; padding: 2px 5px; border-radius: 3px; font-family: Consolas, monospace; font-size: 13px; }
pre { background: #1a1a2e; color: #e0e0e0; padding: 14px; border-radius: 6px; overflow-x: auto; font-size: 13px; }
pre code { background: none; color: inherit; padding: 0; }
strong { color: #0721A8; }
hr { display: none; }
ul, ol { padding-left: 24px; }
li { margin-bottom: 4px; }
img { max-width: 100%; height: auto; display: block; margin: 16px auto; }
a { color: #0D47A1; text-decoration: none; }
"""

EDGE_PATHS = [
    r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files\Google\Chrome\Application\chrome.exe",
]


def find_browser():
    for p in EDGE_PATHS:
        if os.path.exists(p):
            return p
    for name in ("msedge", "chrome", "chromium"):
        found = shutil.which(name)
        if found:
            return found
    raise FileNotFoundError("找不到 Edge 或 Chrome")


def read_md(md_file: Path) -> str:
    content = md_file.read_text(encoding="utf-8")
    if content.startswith("---"):
        end = content.find("---", 3)
        if end != -1:
            content = content[end + 3:].strip()
    return content


# ============ PDF 导出 ============
def to_pdf(md_file: Path, out_file: Path):
    import markdown
    content = read_md(md_file)
    body = markdown.markdown(content, extensions=["tables", "fenced_code", "toc"])

    # 内嵌图片为 base64
    def inline_img(m):
        src = m.group(1)
        p = md_file.parent / src
        if p.exists():
            ext = p.suffix.lower().lstrip(".")
            mime = "image/svg+xml" if ext == "svg" else f"image/{ext}"
            data = base64.b64encode(p.read_bytes()).decode()
            return f'src="data:{mime};base64,{data}"'
        return m.group(0)

    body = re.sub(r'src="([^"]+)"', inline_img, body)
    html = f'<!DOCTYPE html><html><head><meta charset="utf-8"><style>{CSS}</style></head><body>{body}</body></html>'

    tmp = Path(tempfile.mkdtemp(prefix="mdexp_"))
    html_file = tmp / "doc.html"
    html_file.write_text(html, encoding="utf-8")
    if out_file.exists():
        out_file.unlink()

    subprocess.run([
        find_browser(), "--headless=new", "--disable-gpu", "--no-pdf-header-footer",
        "--run-all-compositor-stages-before-draw", "--virtual-time-budget=15000",
        f"--print-to-pdf={out_file}", html_file.as_uri(),
    ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=90)

    for _ in range(30):
        if out_file.exists() and out_file.stat().st_size > 1024:
            break
        time.sleep(0.5)
    shutil.rmtree(tmp, ignore_errors=True)
    if not out_file.exists() or out_file.stat().st_size <= 1024:
        raise RuntimeError("PDF 生成失败")


# ============ Word 导出 ============
def to_word(md_file: Path, out_file: Path):
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    content = read_md(md_file)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 设置内置 Heading 样式的字体和颜色（这样导航窗格能识别）
    # 全黑配色：标题/加粗/目录统一黑色（原 360 品牌蓝已停用）
    BLUE_DARK = RGBColor(0x00, 0x00, 0x00)
    BLUE_MID = RGBColor(0x00, 0x00, 0x00)
    YH = "Microsoft YaHei"

    def setup_heading_style(style_name, size, color):
        """配置内置标题样式的字体"""
        style = doc.styles[style_name]
        style.font.name = YH
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), YH)

    setup_heading_style("Heading 1", 22, BLUE_DARK)
    setup_heading_style("Heading 2", 16, BLUE_DARK)
    setup_heading_style("Heading 3", 14, BLUE_MID)
    setup_heading_style("Heading 4", 12, BLUE_MID)

    def font(run, size=12, bold=False, color=None, name="Microsoft YaHei"):
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), name)

    def heading(text, level):
        """使用 Word 内置 Heading 样式（导航窗格+TOC 可识别）"""
        style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
        p = doc.add_paragraph(style=style_map.get(level, "Heading 4"))
        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        font(run, [22, 16, 14, 12][min(level, 4) - 1], True,
             BLUE_DARK if level <= 2 else BLUE_MID)
        return p

    def add_toc():
        """插入 Word TOC 域（打开后右键更新即可跳转）"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run("目  录"), 16, True, BLUE_DARK)
        # 插入 TOC 域
        p2 = doc.add_paragraph()
        run = p2.add_run()
        fldChar1 = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._r.append(fldChar1)
        run2 = p2.add_run()
        instrText = run2._r.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        run2._r.append(instrText)
        run3 = p2.add_run()
        fldChar2 = run3._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "separate"})
        run3._r.append(fldChar2)
        run4 = p2.add_run("（右键此处 → 更新域 → 更新整个目录）")
        font(run4, 10, False, RGBColor(0x99, 0x99, 0x99))
        run5 = p2.add_run()
        fldChar3 = run5._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run5._r.append(fldChar3)
        doc.add_paragraph()  # 空行

    def blockquote(text):
        """引用块：蓝色左边框 + 浅蓝底"""
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        left = pBdr.makeelement(qn("w:left"), {
            qn("w:val"): "single", qn("w:sz"): "12",
            qn("w:space"): "4", qn("w:color"): "0721A8"
        })
        pBdr.append(left)
        pPr.append(pBdr)
        shd = pPr.makeelement(qn("w:shd"), {qn("w:fill"): "E8EEF6", qn("w:val"): "clear"})
        pPr.append(shd)
        for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
            if not part:
                continue
            if i % 2 == 0:
                font(p.add_run(part), 11)
            else:
                font(p.add_run(part), 11, True, RGBColor(0x00, 0x00, 0x00))

    def para(text):
        p = doc.add_paragraph()
        for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
            if not part:
                continue
            if i % 2 == 0:
                font(p.add_run(part), 12)
            else:
                font(p.add_run(part), 12, True, RGBColor(0x00, 0x00, 0x00))

    def table(headers, rows):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = ""
            # 表头：去掉 ** 标记，始终加粗黑字无底色
            h_clean = h.replace("**", "")
            font(c.paragraphs[0].add_run(h_clean), 10, True, RGBColor(0x00, 0x00, 0x00))
            tcPr = c._element.get_or_add_tcPr()
            shd = tcPr.makeelement(qn("w:shd"), {qn("w:fill"): "FFFFFF", qn("w:val"): "clear"})
            tcPr.append(shd)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                if ci < len(headers):
                    c = t.rows[ri + 1].cells[ci]
                    c.text = ""
                    # 单元格内处理加粗
                    for j, part in enumerate(re.split(r"\*\*(.+?)\*\*", val)):
                        if not part:
                            continue
                        if j % 2 == 0:
                            font(c.paragraphs[0].add_run(part), 10)
                        else:
                            font(c.paragraphs[0].add_run(part), 10, True, RGBColor(0x00, 0x00, 0x00))
        doc.add_paragraph()

    lines = content.split("\n")
    buf = []
    in_tbl = False
    in_toc_section = False  # 跳过 md 里手写的"目录"段落
    toc_inserted = False
    org_headings = {}  # heading text -> paragraph, for image insertion

    def flush():
        nonlocal buf
        if not buf:
            return
        hs = [c.strip() for c in buf[0].split("|")[1:-1]]
        rs = [[c.strip() for c in ln.split("|")[1:-1]] for ln in buf[2:]]
        table(hs, rs)
        buf = []

    i = 0
    while i < len(lines):
        ln = lines[i].strip()
        if ln.startswith("!["):
            i += 1
            continue
        # --- 处理引用块 (> 开头) ---
        if ln.startswith(">"):
            quote_text = ln.lstrip(">").strip()
            # 合并连续的 > 行
            while i + 1 < len(lines) and lines[i + 1].strip().startswith(">"):
                i += 1
                next_line = lines[i].strip().lstrip(">").strip()
                if next_line:
                    quote_text += "\n" + next_line
            if quote_text:
                blockquote(quote_text)
            i += 1
            continue
        # --- 处理分隔线 ---
        if ln == "---" or ln == "***" or ln == "___":
            i += 1
            continue
        # --- 处理列表项 (- 或 * 或数字.) ---
        list_m = re.match(r"^[-*]\s+(.*)$", ln) or re.match(r"^\d+\.\s+(.*)$", ln)
        if list_m:
            if in_tbl:
                in_tbl = False
                flush()
            list_text = list_m.group(1)
            p = doc.add_paragraph()
            # 用 bullet 样式的缩进
            pPr = p._element.get_or_add_pPr()
            ind = pPr.makeelement(qn("w:ind"), {qn("w:left"): "480", qn("w:hanging"): "240"})
            pPr.append(ind)
            # 加个 bullet 符号
            bullet_char = "• " if ln.startswith("-") or ln.startswith("*") else re.match(r"^(\d+\.)\s", ln).group(1) + " "
            font(p.add_run(bullet_char), 12)
            for j, part in enumerate(re.split(r"\*\*(.+?)\*\*", list_text)):
                if not part:
                    continue
                if j % 2 == 0:
                    font(p.add_run(part), 12)
                else:
                    font(p.add_run(part), 12, True, RGBColor(0x00, 0x00, 0x00))
            i += 1
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", ln)
        if m:
            if in_tbl:
                in_tbl = False
                flush()
            lvl = len(m.group(1))
            htxt = m.group(2)
            # 检测手写目录段（"## 目录"），跳过该段下面的列表
            if "目录" == htxt.strip():
                in_toc_section = True
                i += 1
                continue
            # 遇到下一个标题，结束 TOC 跳过区
            if in_toc_section:
                in_toc_section = False
            p = heading(htxt, lvl)
            org_headings[htxt] = p
            # 第一个 H1 后插入自动目录
            if lvl == 1 and not toc_inserted:
                toc_inserted = True
                add_toc()
            i += 1
            continue
        # 跳过手写目录段内的列表行
        if in_toc_section:
            i += 1
            continue
        if ln.startswith("|") and ln.count("|") >= 2:
            if not in_tbl:
                in_tbl = True
                buf = []
            buf.append(ln)
            i += 1
            continue
        elif in_tbl:
            in_tbl = False
            flush()
            continue
        if not ln:
            i += 1
            continue
        para(ln)
        i += 1
    if in_tbl:
        flush()

    # 内嵌图片(找 md 里的 ![...](path) 引用),插到对应位置
    img_refs = re.findall(r"!\[.*?\]\(([^)]+)\)", content)
    if img_refs:
        # 简单策略:图片插在第一个含"架构"的标题后,否则文末
        img_para = None
        for htext, hpara in org_headings.items():
            if "架构" in htext or "组织" in htext:
                img_para = hpara
                break
        for src in img_refs:
            ipath = md_file.parent / src
            if not ipath.exists():
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(ipath), width=Inches(6.0))
            if img_para is not None:
                p._element.getparent().remove(p._element)
                img_para._element.addnext(p._element)
                img_para = p  # next image goes after this one

    doc.save(str(out_file))


# ============ 主流程 ============
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("md", help="输入 md 文件")
    ap.add_argument("--to", choices=["pdf", "word", "docx"], default="pdf", help="输出格式")
    ap.add_argument("--out", dest="out", default=None, help="输出路径(可选)")
    args = ap.parse_args()

    fmt = "word" if args.to in ("word", "docx") else "pdf"
    ext = ".docx" if fmt == "word" else ".pdf"

    src = Path(args.md).resolve()
    if not src.exists():
        print(f"[ERROR] 找不到: {src}")
        sys.exit(1)
    dst = Path(args.out).resolve() if args.out else src.with_suffix(ext)

    # 中文路径规避:复制到英文临时目录处理
    work = Path(tempfile.mkdtemp(prefix="mdexp_work_"))
    tmp_md = work / "doc.md"
    shutil.copy2(src, tmp_md)
    # 复制同目录的图片(md 引用的)
    content = src.read_text(encoding="utf-8")
    for imref in re.findall(r"!\[.*?\]\(([^)]+)\)", content):
        ip = src.parent / imref
        if ip.exists():
            (work / imref).parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(ip, work / imref)
    tmp_out = work / ("doc" + ext)

    if fmt == "pdf":
        to_pdf(tmp_md, tmp_out)
    else:
        to_word(tmp_md, tmp_out)

    shutil.copy2(tmp_out, dst)
    shutil.rmtree(work, ignore_errors=True)

    size_kb = round(dst.stat().st_size / 1024)
    print(f"[OK] {dst.name} ({size_kb} KB) [{fmt.upper()}]")


if __name__ == "__main__":
    main()
